from pathlib import Path
import mammoth
from docx import Document


def substituir_texto_docx(doc, variaveis):
    """
    Substitui marcadores no DOCX preservando boa parte da formatação.
    Funciona melhor quando o marcador está inteiro dentro do mesmo run.
    """

    def substituir_em_paragrafos(paragrafos):
        for paragrafo in paragrafos:
            for run in paragrafo.runs:
                for marcador, valor in variaveis.items():
                    if marcador in run.text:
                        run.text = run.text.replace(marcador, str(valor))

    substituir_em_paragrafos(doc.paragraphs)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                substituir_em_paragrafos(celula.paragraphs)

    return doc


def converter_docx_para_html(docx_path, html_path):
    docx_path = Path(docx_path)
    html_path = Path(html_path)

    html_path.parent.mkdir(parents=True, exist_ok=True)

    style_map = """
    p[style-name='Title'] => h1:fresh
    p[style-name='Título'] => h1:fresh
    p[style-name='Heading 1'] => h1:fresh
    p[style-name='Título 1'] => h1:fresh
    p[style-name='Heading 2'] => h2:fresh
    p[style-name='Título 2'] => h2:fresh

    r[style-name='Emphasis'] => em
    r[style-name='Ênfase'] => em
    r[style-name='Intense Emphasis'] => strong > em
    r[style-name='Ênfase Intensa'] => strong > em
    """

    with open(docx_path, "rb") as docx_file:
        resultado = mammoth.convert_to_html(
            docx_file,
            style_map=style_map,
            include_default_style_map=True
        )

    html_body = resultado.value

    if resultado.messages:
        print("Avisos do Mammoth:")
        for msg in resultado.messages:
            print(msg)

    html_completo = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <title>Contrato</title>

    <style>
        @page {{
            size: A4;
            margin: 2.6cm 1.7cm 2.4cm 2.35cm;
        }}

        * {{
            box-sizing: border-box;
        }}

        body {{
            margin: 0;
            padding: 0;
            background: #e5e5e5;
            color: #000;
            font-family: Arial, "Arial MT", sans-serif;
            font-size: 11pt;
            line-height: 1.15;
        }}

        .pagina {{
            width: 21cm;
            min-height: 29.7cm;
            margin: 30px auto;
            padding: 2.6cm 1.7cm 2.4cm 2.35cm;
            background: #fff;
            box-shadow: 0 0 10px rgba(0, 0, 0, 0.18);
        }}

        p {{
            margin: 0 0 10pt 0;
            text-align: justify;
            text-indent: 0;
        }}

        h1 {{
            font-size: 13pt;
            line-height: 1.15;
            text-align: center;
            text-transform: uppercase;
            margin: 0 0 18pt 0;
            font-weight: bold;
        }}

        h2 {{
            font-size: 12pt;
            line-height: 1.15;
            text-align: left;
            text-transform: uppercase;
            margin: 18pt 0 10pt 0;
            font-weight: bold;
        }}

        h3 {{
            font-size: 11pt;
            line-height: 1.15;
            text-align: left;
            margin: 14pt 0 8pt 0;
            font-weight: bold;
        }}

        strong, b {{
            font-weight: bold;
        }}

        em, i {{
            font-style: italic;
        }}

        strong em,
        em strong,
        b i,
        i b {{
            font-weight: bold;
            font-style: italic;
        }}

        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 12pt 0 18pt 0;
            page-break-inside: avoid;
            font-size: 10.5pt;
        }}

        td, th {{
            border: 1px solid #000;
            padding: 5pt 6pt;
            vertical-align: top;
            text-align: justify;
            line-height: 1.15;
        }}

        th {{
            font-weight: bold;
            text-align: center;
        }}

        table tr:first-child td,
        table tr:first-child th {{
            font-weight: bold;
            text-align: center;
            text-transform: uppercase;
        }}

        table td:first-child {{
            width: 34%;
            font-weight: bold;
            text-align: left;
        }}

        table td:nth-child(2) {{
            width: 66%;
        }}

        ol, ul {{
            margin-top: 6pt;
            margin-bottom: 10pt;
            padding-left: 28pt;
        }}

        li {{
            margin-bottom: 6pt;
            text-align: justify;
            line-height: 1.15;
        }}

        ol li p,
        ul li p {{
            margin: 0;
            text-align: justify;
        }}

        h1, h2, h3 {{
            break-after: avoid;
            page-break-after: avoid;
        }}

        @media print {{
            body {{
                background: #fff;
            }}

            .pagina {{
                width: auto;
                min-height: auto;
                margin: 0;
                padding: 0;
                box-shadow: none;
            }}
        }}
    </style>
</head>
<body>
    <div class="pagina">
        {html_body}
    </div>
</body>
</html>
"""

    html_path.write_text(html_completo, encoding="utf-8")

    print("HTML gerado em:", html_path.resolve())

    return html_path


def gerar_contrato_html(dados, output_path):
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    template_path = Path(__file__).resolve().parent / "template_contrato" / "contrato_financiamento_comprador_unico.docx"

    doc = Document(template_path)

    variaveis = dados

    doc = substituir_texto_docx(doc, variaveis)

    docx_saida = output_path.with_suffix(".docx")
    doc.save(docx_saida)

    print("Arquivo .docx gerado em:", docx_saida.resolve())
